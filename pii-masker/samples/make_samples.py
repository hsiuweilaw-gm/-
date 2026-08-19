# -*- coding: utf-8 -*-
"""產生示範用檔案（內容皆為虛構個資），供測試遮罩工具使用。

執行：python samples/make_samples.py [輸出資料夾]
"""
from __future__ import annotations

import sys
from pathlib import Path

FAKE = {
    "name": "王小明",
    "id": "A123456789",           # 經典演算法測試號碼，非真實個資
    "birth": "民國75年3月12日",
    "mobile": "0912-345-678",
    "tel": "(02)2712-3456",
    "email": "hsiao.ming@example.com.tw",
    "addr": "台北市大安區和平東路二段106巷3號5樓",
    "card": "4111-1111-1111-1111",  # 測試卡號
}


def make_docx(out: Path):
    from docx import Document
    doc = Document()
    doc.add_heading("保險理賠申請書（示範文件）", level=1)
    doc.add_paragraph("申請人：%(name)s（身分證字號：%(id)s）" % FAKE)
    doc.add_paragraph("出生日期：%(birth)s" % FAKE)
    doc.add_paragraph("聯絡電話：%(mobile)s／%(tel)s" % FAKE)
    doc.add_paragraph("電子郵件：%(email)s" % FAKE)
    doc.add_paragraph("通訊地址：%(addr)s" % FAKE)
    table = doc.add_table(rows=2, cols=2)
    table.rows[0].cells[0].text = "被保險人"
    table.rows[0].cells[1].text = FAKE["name"]
    table.rows[1].cells[0].text = "信用卡號"
    table.rows[1].cells[1].text = FAKE["card"]
    doc.save(out / "示範_理賠申請書.docx")


def make_xlsx(out: Path):
    from openpyxl import Workbook
    wb = Workbook()
    ws = wb.active
    ws.title = "客戶名單"
    ws.append(["姓名", "身分證字號", "行動電話", "電子郵件", "地址"])
    ws.append([FAKE["name"], FAKE["id"], FAKE["mobile"], FAKE["email"], FAKE["addr"]])
    ws.append(["要保人：李大華", "統一編號：04595257", "電話 0987654321", "", ""])
    wb.save(out / "示範_客戶名單.xlsx")


def make_pdf(out: Path):
    import pymupdf
    doc = pymupdf.open()
    page = doc.new_page()
    lines = [
        "保險契約通知書（示範文件）",
        "",
        "要保人：%(name)s" % FAKE,
        "身分證字號：%(id)s" % FAKE,
        "行動電話：%(mobile)s" % FAKE,
        "電子郵件：%(email)s" % FAKE,
        "通訊地址：%(addr)s" % FAKE,
        "出生日期：%(birth)s" % FAKE,
    ]
    page.insert_text((72, 90), "\n".join(lines),
                     fontname="china-t", fontsize=13, lineheight=1.6)
    doc.save(out / "示範_契約通知書.pdf")
    doc.close()


def main():
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else Path(__file__).parent
    out.mkdir(parents=True, exist_ok=True)
    make_docx(out)
    make_xlsx(out)
    make_pdf(out)
    print("示範檔案已產生於：%s" % out.resolve())


if __name__ == "__main__":
    main()
