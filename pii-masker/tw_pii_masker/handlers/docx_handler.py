# -*- coding: utf-8 -*-
"""Word (.docx) 處理器。

逐段掃描本文、表格（含巢狀表格）、各節頁首/頁尾，於 run 層級原地取代，
盡可能保留原始格式；同時預設清除文件屬性中的作者資訊。

已知限制：
  - 舊版 .doc 不支援（請先以 Word 另存為 .docx）
  - 文字方塊（text box）、SmartArt 內的文字無法處理
  - 個資若被「分散在格式不同的多個 run」中，取代後該段格式以第一個 run 為準
"""
from __future__ import annotations

from typing import Iterable, Iterator, List, Tuple

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

from ..engine import MaskedItem, MaskingEngine
from ..report import Report


def _replace_span(runs, start: int, end: int, replacement: str) -> None:
    """把 runs 串接文字中 [start, end) 的區段換成 replacement。

    取代文字放入第一個重疊的 run，其餘重疊部分刪除，保留各 run 格式。
    """
    pos = 0
    bounds: List[Tuple[int, int, object]] = []
    for r in runs:
        length = len(r.text)
        bounds.append((pos, pos + length, r))
        pos += length
    first = True
    for s, e, r in bounds:
        if e <= start or s >= end:
            continue
        local_start = max(start - s, 0)
        local_end = min(end - s, e - s)
        t = r.text
        if first:
            r.text = t[:local_start] + replacement + t[local_end:]
            first = False
        else:
            r.text = t[:local_start] + t[local_end:]


def _mask_paragraph(para: Paragraph, engine: MaskingEngine,
                    report: Report, location: str, hint: str = "") -> int:
    runs = para.runs
    text = "".join(r.text for r in runs)
    if not text:
        return 0
    matches = engine.scan(text, context_hint=hint)
    if not matches:
        return 0
    items: List[MaskedItem] = []
    # 由後往前取代，前面的位移才不會失效
    for m in reversed(matches):
        repl = engine.replacement(m)
        _replace_span(runs, m.start, m.end, repl)
        items.append(MaskedItem(m.type, m.label, m.text, repl))
    items.reverse()
    report.add_all(location, items)
    return len(items)


def _iter_block_paragraphs(parent, prefix: str) -> Iterator[Tuple[Paragraph, str, str]]:
    """走訪一個容器（文件本文 / 儲存格 / 頁首頁尾）內的段落與表格。

    產出 (段落, 位置描述, 前後文提示)。表格儲存格的提示取自
    「左側儲存格」與「同欄第一列（表頭）」，讓「標籤在左欄、值在右欄」
    這類表單也能正確判斷姓名／出生日期等需要前後文的類型。
    """
    for i, para in enumerate(parent.paragraphs, 1):
        yield para, "%s 段落%d" % (prefix, i), ""
    for ti, table in enumerate(parent.tables, 1):
        yield from _iter_table(table, "%s 表格%d" % (prefix, ti))


def _iter_table(table: Table, prefix: str) -> Iterator[Tuple[Paragraph, str, str]]:
    for ri, row in enumerate(table.rows, 1):
        for ci, cell in enumerate(row.cells, 1):
            cell_prefix = "%s(%d,%d)" % (prefix, ri, ci)
            hints = []
            if ci >= 2:
                hints.append(row.cells[ci - 2].text)      # 左側儲存格
            if ri >= 2 and ci <= len(table.rows[0].cells):
                hints.append(table.rows[0].cells[ci - 1].text)  # 同欄表頭
            hint = "\n".join(h for h in hints if h)
            for para in cell.paragraphs:
                yield para, cell_prefix, hint
            for ti, sub in enumerate(cell.tables, 1):
                yield from _iter_table(sub, "%s 巢狀表格%d" % (cell_prefix, ti))


def mask_docx(input_path: str, output_path: str,
              engine: MaskingEngine, report: Report,
              keep_metadata: bool = False) -> None:
    doc = Document(input_path)

    for para, location, hint in _iter_block_paragraphs(doc, "本文"):
        _mask_paragraph(para, engine, report, location, hint)

    for si, section in enumerate(doc.sections, 1):
        parts = (
            (section.header, "頁首"),
            (section.footer, "頁尾"),
            (section.first_page_header, "首頁頁首"),
            (section.first_page_footer, "首頁頁尾"),
            (section.even_page_header, "偶數頁頁首"),
            (section.even_page_footer, "偶數頁頁尾"),
        )
        for part, name in parts:
            if part is None:
                continue
            prefix = "第%d節%s" % (si, name)
            for para, location, hint in _iter_block_paragraphs(part, prefix):
                _mask_paragraph(para, engine, report, location, hint)

    if not keep_metadata:
        props = doc.core_properties
        if props.author or props.last_modified_by:
            report.warn("已清除文件屬性中的作者／最後修改者資訊")
        props.author = ""
        props.last_modified_by = ""
        props.comments = ""

    doc.save(output_path)
